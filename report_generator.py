import os
import re
import xml.sax.saxutils as saxutils
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from datetime import datetime


def _pdf_text(value):
    """
    Make a value safe to drop into a ReportLab Paragraph.

    Paragraph() doesn't render plain text — it parses a small, strict
    markup language and raises ValueError on anything outside it (a bare
    '<', an unescaped '&', a non-self-closing tag like '<br>' instead of
    '<br/>', etc). The Groq-generated AI insights routinely contain
    markdown (**bold**, '<br>' line breaks) that trips this up and crashes
    PDF generation entirely.

    This escapes everything as XML first (so nothing in the source text
    can be misread as a tag), then re-introduces just the handful of tags
    ReportLab actually supports.
    """
    if value is None:
        return ""

    text = saxutils.escape(str(value))
    text = re.sub(r"&lt;br\s*/?&gt;", "<br/>", text, flags=re.IGNORECASE)
    text = text.replace("\n", "<br/>")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    return text


def _get_model_metric(model_result):
    metric_name = model_result.get("metric_name", "Score")
    metric_value = model_result.get("metric_value", model_result.get("accuracy", model_result.get("score")))
    return metric_name, metric_value

def generate_pdf(summary, stats, insights, model_result, chart_paths, ai_insights):
    os.makedirs("outputs/reports", exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pdf_path = f"outputs/reports/report_{timestamp}.pdf"
    doc = SimpleDocTemplate(pdf_path)
    styles = getSampleStyleSheet()

    elements = []

    #Title
    elements.append(Paragraph("NestInsight Report", styles['Title']))
    elements.append(Spacer(1, 20))
    #Summary
    elements.append(Paragraph("Dataset Summary:", styles['Heading2']))
    for key, value in summary.items():
        elements.append(Paragraph(f"{_pdf_text(key)}: {_pdf_text(value)}", styles['BodyText']))
    elements.append(Spacer(1, 20))

    #Statistics
    elements.append(Paragraph("Statistical Analysis:", styles['Heading2']))
    for column, value in stats.items():
        elements.append(Paragraph(f"<b>{_pdf_text(column)}</b>", styles['BodyText']))
        elements.append(Paragraph(_pdf_text(value), styles['BodyText']))
    elements.append(Spacer(1, 10))

    # Business Insights
    elements.append(Paragraph("Business Insights:", styles['Heading2']))

    for insight in insights:
        elements.append(
            Paragraph(
                _pdf_text(f"• {insight}"),
                styles['BodyText']
            )
        )

    elements.append(Spacer(1, 20))

    # AI Business Intelligence
    elements.append(
        Paragraph(
            "AI Business Intelligence:",
            styles['Heading2']
        )
    )

    # Split AI insights line by line for clean formatting
    for line in ai_insights.split("\n"):
        line = line.strip()

        if not line:
            continue

        # Section headings
        if line.endswith(":"):
            elements.append(
                Paragraph(
                    _pdf_text(line),
                    styles['Heading3']
                )
            )

        # Bullet points
        elif line.startswith("-"):
            elements.append(
                Paragraph(
                    _pdf_text(f"• {line[1:].strip()}"),
                    styles['BodyText']
                )
            )

        # Normal text fallback
        else:
            elements.append(
                Paragraph(
                    _pdf_text(line),
                    styles['BodyText']
                )
            )

    elements.append(Spacer(1, 20))

    #Model Results
    elements.append(Paragraph("Model Results:", styles['Heading2']))
    
    if model_result.get("status") == "success":
        metric_name, metric_value = _get_model_metric(model_result)
        suffix = "%" if metric_name == "Accuracy" else ""
        elements.append(Paragraph(_pdf_text(f"{metric_name}: {metric_value}{suffix}"), styles['BodyText']))
        elements.append(Paragraph(_pdf_text(f"Model Path: {model_result['model_path']}"), styles['BodyText']))
    else:
        elements.append(Paragraph(_pdf_text(f"Training Failed: {model_result['message']}"), styles['BodyText']))

    elements.append(Spacer(1, 20))

    #Charts
    elements.append(Paragraph("Visual Analytics:", styles['Heading2']))
    for chart_name, chart_path in chart_paths.items():
        elements.append(Paragraph(chart_name.capitalize(), styles['Heading3']))
        elements.append(Image(chart_path, width=400, height=250))
        elements.append(Spacer(1, 20))

    doc.build(elements)

    return pdf_path

# Generate DOCX report
def generate_docx(summary, stats, insights, model_result, chart_paths, ai_insights):
    os.makedirs("outputs/reports", exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")    
    docx_path = f"outputs/reports/report_{timestamp}.docx"
    doc = Document()
    #Title
    doc.add_heading("NestInsight Report", 0)

    #Summary
    doc.add_heading("Dataset Summary:", level=1)
    for key, value in summary.items():
        doc.add_paragraph(f"{key}: {value}")

    #Statistics
    doc.add_heading("Statistical Analysis:", level=1)
    for column, values in stats.items():
        doc.add_paragraph(f"{column}: {values}")
        
    #Insights
    doc.add_heading("Business Insights:", level=1)
    for insight in insights:
        doc.add_paragraph(insight)

    doc.add_heading("AI Business Intelligence", level=1)
    doc.add_paragraph(ai_insights)
    #Model Results
    doc.add_heading("Model Results:", level=1)

    if model_result.get("status") == "success":
        metric_name, metric_value = _get_model_metric(model_result)
        suffix = "%" if metric_name == "Accuracy" else ""
        doc.add_paragraph(f"{metric_name}: {metric_value}{suffix}")
        doc.add_paragraph(f"Model Path: {model_result['model_path']}")
    else:
        doc.add_paragraph(f"Training Failed: {model_result['message']}")

    #Charts
    doc.add_heading("Visual Analytics:", level=1)
    for chart_name, chart_path in chart_paths.items():
        doc.add_heading(chart_name.capitalize(), level=2)
        doc.add_picture(chart_path, width=None)
    doc.save(docx_path)
    return docx_path